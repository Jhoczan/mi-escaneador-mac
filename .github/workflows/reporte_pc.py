import platform
import subprocess
import os
import psutil
from docx import Document

# Intentar importar WMI solo en Windows
try:
    if platform.system() == "Windows":
        import wmi
except ImportError:
    wmi = None

def abrir_archivo(ruta):
    try:
        if platform.system() == "Windows":
            os.startfile(ruta)
        elif platform.system() == "Darwin":
            subprocess.call(["open", ruta])
    except:
        pass

def obtener_datos():
    sistema_act = platform.system()
    res = {}
    serial = "Generico"

    if sistema_act == "Darwin": # MAC
        def run(cmd): return subprocess.check_output(cmd, shell=True).decode().strip()
        serial = run("system_profiler SPHardwareDataType | awk '/Serial Number/ {print $4}'")
        res = {
            "S.O.": f"macOS {run('sw_vers -productVersion')}",
            "Serie": serial,
            "Modelo": run("system_profiler SPHardwareDataType | awk '/Model Name/ {print $3,$4}'"),
            "RAM": f"{round(psutil.virtual_memory().total / (1024**3), 2)} GB",
            "Office": "Instalado" if os.path.exists("/Applications/Microsoft Word.app") else "No detectado"
        }
    else: # WINDOWS
        import wmi
        c = wmi.WMI()
        win = c.Win32_OperatingSystem()[0]
        bios = c.Win32_BIOS()[0]
        serial = bios.SerialNumber.strip()
        res = {
            "S.O.": win.Caption,
            "Serie": serial,
            "Modelo": c.Win32_ComputerSystem()[0].Model,
            "RAM": f"{round(int(c.Win32_ComputerSystem()[0].TotalPhysicalMemory) / (1024**3), 2)} GB",
            "Office": "Detectado" if os.path.exists(r"C:\Program Files\Microsoft Office") else "No detectado"
        }
    return res, serial

def ejecutar():
    print("=== GENERADOR DE REPORTE ===")
    nombre_usr = input("Nombre del archivo (Enter para usar Serial): ").strip()
    datos, sn = obtener_datos()
    
    archivo = nombre_usr if nombre_usr else f"Reporte_{sn}"
    if not archivo.endswith(".docx"): archivo += ".docx"

    doc = Document()
    doc.add_heading('Ficha Técnica del Equipo', 0)
    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = 'Table Grid'
    for k, v in datos.items():
        celdas = tabla.add_row().cells
        celdas[0].text, celdas[1].text = str(k), str(v)
    
    doc.save(archivo)
    print(f"Listo: {archivo}")
    abrir_archivo(archivo)

if __name__ == "__main__":
    ejecutar()
